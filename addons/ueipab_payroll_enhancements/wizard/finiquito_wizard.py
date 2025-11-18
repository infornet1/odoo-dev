# -*- coding: utf-8 -*-

from odoo import models, fields, api
from odoo.exceptions import UserError
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io


class FiniquitoWizard(models.TransientModel):
    _name = 'finiquito.wizard'
    _description = 'Acuerdo Finiquito Laboral Wizard'

    payslip_ids = fields.Many2many(
        'hr.payslip',
        string='Payslips',
        required=True,
        domain="[('struct_id.name', 'in', ['Liquidación Venezolana', 'Liquidación Venezolana V2']), ('state', '!=', 'cancel')]",
        help='Select liquidation payslips for finiquito agreement'
    )

    currency_id = fields.Many2one(
        'res.currency',
        string='Currency',
        required=True,
        default=lambda self: self.env.ref('base.USD'),
        help='Currency for displaying amounts'
    )

    output_format = fields.Selection(
        [('pdf', 'PDF'), ('docx', 'Word Document (.docx)')],
        string='Output Format',
        default='pdf',
        required=True,
        help='Choose output format for the report'
    )

    # Exchange Rate Override Options (added in v1.25.0)
    use_custom_rate = fields.Boolean(
        string='Use Custom Exchange Rate',
        default=False,
        help='Override automatic exchange rate lookup with a custom rate'
    )

    custom_exchange_rate = fields.Float(
        string='Custom Rate (VEB/USD)',
        digits=(12, 4),
        help='Custom exchange rate to use (e.g., 236.4601). Only used if "Use Custom Exchange Rate" is checked.'
    )

    rate_date = fields.Date(
        string='Rate Date',
        help='Date to lookup exchange rate automatically. Leave blank to use payslip liquidation date.'
    )

    def action_generate_report(self):
        """Generate report in selected format"""
        self.ensure_one()

        if self.output_format == 'pdf':
            return self.action_print_pdf()
        else:
            return self.action_export_docx()

    def action_print_pdf(self):
        """Generate PDF report"""
        self.ensure_one()

        data = {
            'payslip_ids': self.payslip_ids.ids,
            'currency_id': self.currency_id.id,
            'use_custom_rate': self.use_custom_rate,
            'custom_exchange_rate': self.custom_exchange_rate,
            'rate_date': self.rate_date,
        }

        return self.env.ref('ueipab_payroll_enhancements.action_report_finiquito').report_action(
            self.payslip_ids, data=data
        )

    def action_export_docx(self):
        """Generate DOCX report"""
        self.ensure_one()

        if not self.payslip_ids:
            raise UserError('Please select at least one payslip.')

        # Get report model to fetch data
        report_model = self.env['report.ueipab_payroll_enhancements.finiquito_report']
        data = {
            'payslip_ids': self.payslip_ids.ids,
            'currency_id': self.currency_id.id,
            'use_custom_rate': self.use_custom_rate,
            'custom_exchange_rate': self.custom_exchange_rate,
            'rate_date': self.rate_date,
        }
        report_values = report_model._get_report_values(self.payslip_ids.ids, data=data)

        # Create DOCX document
        doc = Document()

        # Set margins (1.5cm all sides for formal letter)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # Generate document for each payslip
        for idx, report in enumerate(report_values['reports']):
            if idx > 0:
                doc.add_page_break()

            self._add_finiquito_content(doc, report)

        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)

        # Generate filename
        if len(self.payslip_ids) == 1:
            filename = f'Finiquito_{self.payslip_ids[0].employee_id.name}_{self.payslip_ids[0].number}.docx'
        else:
            filename = f'Finiquito_Multiple_{len(self.payslip_ids)}_employees.docx'

        # Return file download
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(docx_file.read()),
            'store_fname': filename,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }

    def _add_finiquito_content(self, doc, report):
        """Add finiquito content to DOCX document"""

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(
            f'FINIQUITO DE LA RELACIÓN LABORAL ENTRE\n'
            f'UNIDAD EDUCATIVA INSTITUTO PRIVADO ANDRES BELLO C.A,\n'
            f'Y {report["employee"].name.upper()}'
        )
        title_run.bold = True
        title_run.font.size = Pt(11)

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run('PAGO DE PRESTACIONES SOCIALES POR TERMINACIÓN DE CONTRATO DE TRABAJO')
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(9)

        doc.add_paragraph()  # Spacing

        # Introduction
        intro = doc.add_paragraph()
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_text = (
            f'Entre: La Unidad Educativa Instituto Privado "Andres Bello" C.A, con Registro de '
            f'Información Fiscal Nro. J-080086171, de este domicilio, debidamente representado por su '
            f'Director Principal GUSTAVO PERDOMO, Venezolano, mayor de edad, de este domicilio y '
            f'titular de cédula de identidad Nro. V-15.128.008, por una parte; y por la otra el trabajador '
            f'{report["employee"].name}, venezolano, mayor de edad, de este domicilio y titular de la cédula '
            f'de identidad Nro. {report["employee"].identification_id or "N/A"}, hemos convenido en finiquitar '
            f'la relación de trabajo, en los siguientes términos:'
        )
        intro.add_run(intro_text).font.size = Pt(9)

        # PRIMERO
        primero = doc.add_paragraph()
        primero.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        primero_run = primero.add_run('PRIMERO: ')
        primero_run.bold = True
        primero_run.font.size = Pt(9)
        primero.add_run(
            f'El Trabajador ha prestado sus servicios en la Instituto Privado Andrés Bello CA, '
            f'desde el {report["date_start_str"]} hasta el {report["date_to_str"]}.'
        ).font.size = Pt(9)

        # SEGUNDO
        segundo = doc.add_paragraph()
        segundo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        segundo_run = segundo.add_run('SEGUNDO: ')
        segundo_run.bold = True
        segundo_run.font.size = Pt(9)
        segundo.add_run(
            f'El Trabajador recibe en este acto, a su entera satisfacción de parte de la Instituto Privado '
            f'Andrés Bello CA, la cantidad de {report["currency"].symbol} {report["net_amount"]:,.2f}, '
            f'suma esta que comprende: todos los conceptos laborales descritos en el recibo de liquidación '
            f'final del contrato de trabajo adjunto y que forma parte de este finiquito.'
        ).font.size = Pt(9)

        # TERCERO
        tercero = doc.add_paragraph()
        tercero.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        tercero_run = tercero.add_run('TERCERO: ')
        tercero_run.bold = True
        tercero_run.font.size = Pt(9)
        tercero.add_run(
            f'El Trabajador {report["employee"].name}, declara que durante todo el tiempo que prestó servicios '
            f'como profesora a la Instituto Privado Andrés Bello CA, recibió oportunamente el pago total de las '
            f'remuneraciones y beneficios convenidos, y en consecuencia nada le adeuda la Instituto Privado Andrés '
            f'Bello CA, por los conceptos indicados en el finiquito ni por ningún otro, sea de origen legal, '
            f'contractual o extracontractual derivado o relacionado con la prestación de sus servicios. Por este '
            f'motivo, la trabajadora otorga a la empresa el más amplio y total finiquito, renunciando a toda acción '
            f'legal, contractual y extracontractual en contra de la Instituto Privado Andrés Bello CA, declaración '
            f'que la trabajadora fórmula libre y espontáneamente en perfecto y cabal conocimiento de cada uno y de '
            f'todos sus derechos.'
        ).font.size = Pt(9)

        # CUARTO
        cuarto = doc.add_paragraph()
        cuarto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cuarto_run = cuarto.add_run('CUARTO: ')
        cuarto_run.bold = True
        cuarto_run.font.size = Pt(9)
        cuarto.add_run(
            'Para constancia firman las partes el presente finiquito en señal de conformidad y se hacen '
            'dos (2) ejemplares del mismo tenor y a un solo efecto.'
        ).font.size = Pt(9)

        doc.add_paragraph()  # Spacing

        # Date line
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.add_run(
            f'El Tigre, a los {report["today_day"]} días del mes de {report["today_month"]} del año {report["today_year"]}.'
        ).font.size = Pt(9)

        # Add extra spacing before signatures
        doc.add_paragraph()
        doc.add_paragraph()

        # Signatures section (using table for layout)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Company signature
        cell_company = table.rows[0].cells[0]
        company_para = cell_company.paragraphs[0]
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_para.add_run('_' * 40 + '\n').font.size = Pt(9)
        company_run1 = company_para.add_run('Instituto Privado Andrés Bello CA\n')
        company_run1.bold = True
        company_run1.font.size = Pt(9)
        company_para.add_run('GUSTAVO PERDOMO\n').font.size = Pt(9)
        company_para.add_run('Director Principal\n').font.size = Pt(9)
        company_para.add_run('C.I.: V-15.128.008').font.size = Pt(9)

        # Employee signature
        cell_employee = table.rows[0].cells[1]
        employee_para = cell_employee.paragraphs[0]
        employee_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        employee_para.add_run('_' * 40 + '\n').font.size = Pt(9)
        employee_run1 = employee_para.add_run(f'TRABAJADOR: {report["employee"].name.upper()}\n')
        employee_run1.bold = True
        employee_run1.font.size = Pt(9)
        employee_para.add_run(
            f'CÉDULA DE IDENTIDAD: {report["employee"].identification_id or "N/A"}'
        ).font.size = Pt(9)
